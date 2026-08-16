import os

def parse_docx(file_path: str) -> str:
    """
    Extracts readable text from a DOCX file using python-docx.
    Includes text from paragraphs and tables.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found at: {file_path}")
        
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        
        # Paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                full_text.append(paragraph.text)
                
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        extracted_text = "\n".join(full_text)
        if not extracted_text.strip():
            raise ValueError("DOCX file contains no extractable text.")
            
        return extracted_text
    except Exception as e:
        raise RuntimeError(f"Error extracting DOCX text: {str(e)}")
